import streamlit as st
import re
import zipfile
from docx import Document
from io import BytesIO
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree


# ══════════════════════════════════════════════════════════════════════════════
# Yardımcı fonksiyonlar
# ══════════════════════════════════════════════════════════════════════════════

def turkish_upper(text):
    replace_map = {"i": "İ", "ı": "I", "ç": "Ç", "ğ": "Ğ",
                   "ö": "Ö", "ş": "Ş", "ü": "Ü"}
    for lower, upper in replace_map.items():
        text = text.replace(lower, upper)
    return text.upper().strip()


def get_effective_spacing(para):
    """
    Paragrafın gerçek satır aralığını döndürür.
    Paragraf düzeyinde ayarlanmamışsa, stil hiyerarşisini yukarı doğru
    tarayarak kalıtılan (inherited) değeri bulur.
    """
    spacing = para.paragraph_format.line_spacing
    if spacing is not None:
        return spacing
    style_obj = para.style
    while style_obj is not None:
        s = style_obj.paragraph_format.line_spacing
        if s is not None:
            return s
        style_obj = style_obj.base_style
    return None


# ══════════════════════════════════════════════════════════════════════════════
# Sabitler
# ══════════════════════════════════════════════════════════════════════════════

# Asla satır aralığı denetimine tabi tutulmayan stiller
MUAF_STILLER = {
    "fbe_kapak_genel",
    "toc 1", "toc 2", "toc 3", "toc 4", "toc 7", "toc 8",
    "table of figures",
    "Caption",            # Şekil/Tablo/Resim başlıkları — zaten 1.0, muaf
    "annotation text",    # Tablo notları — zaten 1.0, muaf
    "fbe_metin_normal_kalın",  # "Sayfa" etiketi satırları
    "fbe_başlık1_sol",
    "List Paragraph",     # Madde/liste işareti paragrafları — kural kapsamı dışı
}

# Bölüm başlık stilleri — modu günceller, kendileri denetlenmez
HEADING_STILLER = {f"Heading {i}" for i in range(1, 10)}

# Bu stiller bulundukları bölümden bağımsız olarak TEK SATIR gerektirir
TEK_SATIR_STILLER = {
    "fbe_metin_sıkışık",   # Önsöz / Teşekkür gövde metni
    # NOT: "Normal (Web)" stili bazı tezlerde ana metin stili olarak kullanıldığından
    # burada sabit kodlanmaz; mod tabanlı kontrol ile ele alınır.
}

# Bu stiller ANA METİN kapsamında değerlendirilir (1.5 aralık)
ANA_METIN_STILLER = {
    "Tez Metni",
    "fbe_metin_normal",    # Kapak jüri listesi + bazı metin blokları
    "Default",             # Etik Beyan metni
    "Normal",              # Kaynakça + genel metin
    "Normal (Web)",        # Bazı tezlerde SONUÇLAR/ABSTRACT gibi ana metin stili
    "Bibliography",        # Kaynakça girişleri (Word'ün yerleşik kaynakça stili)
    "numarasiz",           # Numarasız liste / metin stili
    "Decimal Aligned",     # Ondalık hizalı tablo/metin stili
}

# Tek satır moduna geçiren bölüm başlıkları (metin olarak)
ON_SAYFALAR_TEK = {
    "ÖZET", "ABSTRACT", "ETİK BEYAN", "ÖNSÖZ",
    "TEŞEKKÜR", "KABUL VE ONAY", "DESTEKLEYICI KURULUŞ BEYAN SAYFASI",
}

# Tek satır moduna geçiren liste/dizin başlıkları (metin olarak)
LISTE_BASLIKLAR = {
    "İÇİNDEKİLER",
    "ŞEKİL LİSTESİ",
    "TABLO LİSTESİ",
    "RESİM LİSTESİ",
    "SEMBOL LİSTESİ",
    "KISALTMALAR LİSTESİ",
    "SİMGELER VE KISALTMALAR LİSTESİ",
    "SEMBOL VE KISALTMA LİSTESİ",
    # NOT: ÖZGEÇMİŞ bazı tezlerde 1.5 aralıkla yazıldığından buraya dahil edilmez.
}

# Kaynakça başlıkları
KAYNAKCA_BASLIKLAR = {
    "KAYNAKLAR", "KAYNAKÇA", "REFERENCES", "KAYNAKLAR DİZİNİ",
}


# ══════════════════════════════════════════════════════════════════════════════
# Sayfa numarası denetim yardımcıları
# ══════════════════════════════════════════════════════════════════════════════

def _footer_dosyalarini_analiz_et(docx_stream):
    """
    .docx içindeki tüm footer XML dosyalarını okuyarak her birinin
    sayfa numarası alanı, hizalama, font boyutu ve parantez/çizgi
    durumunu döndürür.
    """
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    if isinstance(docx_stream, str):
        zf = zipfile.ZipFile(docx_stream)
    else:
        docx_stream.seek(0)
        zf = zipfile.ZipFile(docx_stream)

    with zf:
        rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
        rels_root = etree.fromstring(rels_xml.encode())
        rid_to_file = {
            r.get('Id'): r.get('Target')
            for r in rels_root
            if 'footer' in r.get('Type', '').lower()
        }

        footer_data = {}
        for rid, fname in rid_to_file.items():
            try:
                raw = zf.read(f'word/{fname}').decode('utf-8')
                tree = etree.fromstring(raw.encode())

                instr_els = tree.findall(f'.//{{{ns}}}instrText')
                instr_texts = [el.text for el in instr_els if el.text]
                has_page = any('PAGE' in t.upper() for t in instr_texts)

                jc_els = tree.findall(f'.//{{{ns}}}jc')
                alignment = jc_els[0].get(f'{{{ns}}}val') if jc_els else None

                sz_els = tree.findall(f'.//{{{ns}}}sz')
                font_sizes = sorted(set(
                    int(el.get(f'{{{ns}}}val', 0)) / 2
                    for el in sz_els
                    if el.get(f'{{{ns}}}val')
                ))

                # PAGE alanı DIŞINDA kalan metin (parantez/çizgi kontrolü)
                non_field_chars = []
                in_field = False
                for el in tree.iter():
                    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                    if tag == 'fldChar':
                        ftype = el.get(f'{{{ns}}}fldCharType', '')
                        in_field = ftype in ('begin', 'separate')
                        if ftype == 'end':
                            in_field = False
                    elif tag == 't' and not in_field and el.text:
                        non_field_chars.append(el.text)

                non_field_text = ''.join(non_field_chars)
                has_bracket = any(c in non_field_text for c in ['(', ')', '-', '–', '—', '[', ']', '/'])

                footer_data[fname] = {
                    'has_page':       has_page,
                    'alignment':      alignment,
                    'font_sizes':     font_sizes,
                    'instr_texts':    instr_texts,
                    'has_bracket':    has_bracket,
                    'non_field_text': non_field_text,
                }
            except Exception as exc:
                footer_data[fname] = {'error': str(exc)}

    return rid_to_file, footer_data


def sayfa_numarasi_denetle(doc, docx_stream):
    """
    Kural listesi:
      1. Sayfa numarası altta ve ortalanmış olmalı (jc = center).
      2. Yazı boyutu 10 pt olmalı.
      3. Ön sayfalar → küçük harf Romen rakamı (lowerRoman).
      4. Giriş ve sonrası → Arap rakamı (decimal / None).
      5. Parantez veya çizgi kullanılmamalı.
    """
    hatalar = []

    rid_to_file, footer_data = _footer_dosyalarini_analiz_et(docx_stream)

    # Section → footer dosyası eşlemesi (yalnızca 'default' tipi)
    section_footer_map = []
    for si, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        fmt   = pgNumType.get(qn('w:fmt'))   if pgNumType is not None else None
        start = pgNumType.get(qn('w:start')) if pgNumType is not None else None

        # ── Kapak / ara geçiş section'larını muaf tut ──────────────────
        # Kriter: pgNumType.start tanımlı değilse bu section'da sayfa
        # numaralandırması başlatılmamıştır (kapak, ara geçiş sayfaları vb.).
        # Sayfa numarası olmaması bu durumlarda normaldir.
        if start is None:
            continue   # Numaralandırma başlatılmamış section — denetleme

        footer_refs_all = sectPr.findall(qn('w:footerReference'))

        for ref in footer_refs_all:
            ftype = ref.get(qn('w:type'))
            rid   = ref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            fname = rid_to_file.get(rid)
            if fname and ftype == 'default':
                section_footer_map.append((si, fmt, start, fname))

    if not section_footer_map:
        hatalar.append("⚠️ [SAYFA NUMARASI: Footer tanımı bulunamadı — sayfa numarası eksik olabilir!]")
        return hatalar

    # Romen / Arap section varlığını TÜM section'lar üzerinden kontrol et
    # (footer referansı olmayan linked section'lar da sayılır)
    romen_var = False
    arap_var  = False
    for section in doc.sections:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        fmt   = pgNumType.get(qn('w:fmt'))   if pgNumType is not None else None
        start = pgNumType.get(qn('w:start')) if pgNumType is not None else None
        if start is None:
            continue  # numaralandırma başlatılmamış — atla
        if fmt == 'lowerRoman':
            romen_var = True
        else:
            arap_var = True

    if not romen_var:
        hatalar.append(
            "⚠️ [SAYFA NUMARASI: Ön sayfa bölümü (Özet, İçindekiler vb.) için küçük harf Romen "
            "rakamı (i, ii, iii…) bulunamadı! İlgili section'da pgNumType=lowerRoman olmalı.]"
        )
    if not arap_var:
        hatalar.append(
            "⚠️ [SAYFA NUMARASI: Giriş ve sonrası için Arap rakamı (1, 2, 3…) bulunamadı! "
            "İlgili section'da pgNumType=decimal veya tanımsız olmalı.]"
        )

    # Her benzersiz footer dosyasını bir kez denetle
    denetlenen = set()
    for si, fmt, start, fname in section_footer_map:
        if fname in denetlenen:
            continue
        denetlenen.add(fname)

        data = footer_data.get(fname, {})
        if 'error' in data:
            hatalar.append(f"⚠️ [SAYFA NUMARASI: Section {si+1} footer okunamadı: {data['error']}]")
            continue

        etiket = f"Section {si + 1} ({'Romen/ön sayfalar' if fmt == 'lowerRoman' else 'Arap/ana metin'})"

        # Kural 1 — PAGE alanı var mı?
        if not data['has_page']:
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Footer'da sayfa numarası alanı (PAGE) bulunamadı!]"
            )
            continue

        # Kural 2 — Ortalanmış olmalı
        if data['alignment'] != 'center':
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası ortalanmış (alt-orta) olmalı! "
                f"(Algılanan hizalama: {data['alignment'] or 'sol/belirsiz'})]"
            )

        # Kural 3 — 10 pt olmalı
        sizes = data['font_sizes']
        if sizes and not all(abs(s - 10.0) < 0.6 for s in sizes):
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası boyutu 10 pt olmalı! "
                f"(Algılanan: {sizes} pt)]"
            )

        # Kural 4 — Parantez / çizgi olmamalı
        if data['has_bracket']:
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası yanında parantez veya çizgi "
                f"kullanılmamalı! (Algılanan: {repr(data['non_field_text'])})]"
            )

    return hatalar


# ══════════════════════════════════════════════════════════════════════════════
# Ana denetim fonksiyonu
# ══════════════════════════════════════════════════════════════════════════════

def tez_denetle(file_stream):
    # Sayfa numarası denetimi için stream'i iki kez okumamız gerektiğinden kopyala
    file_stream.seek(0)
    docx_bytes = file_stream.read()

    doc = Document(BytesIO(docx_bytes))
    hata_sayisi = 0

    # ── 0. SAYFA NUMARASI DENETİMİ ───────────────────────────────────────
    sayfa_hatalari = sayfa_numarasi_denetle(doc, BytesIO(docx_bytes))
    if sayfa_hatalari:
        # Hataları belgenin ilk paragrafına ekle
        para = doc.paragraphs[0]
        for hata_metni in sayfa_hatalari:
            run = para.add_run("\n" + hata_metni)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
        hata_sayisi += len(sayfa_hatalari)

    # Durum değişkenleri
    kapak_modu = True
    tek_satir_modu = False
    kaynakca_modu = False

    # ── 1. KENAR BOŞLUKLARI ───────────────────────────────────────────────
    for section in doc.sections:
        l = round(section.left_margin.cm, 1)
        r = round(section.right_margin.cm, 1)
        if l != 3.0 or r != 2.5:
            run = doc.paragraphs[0].add_run(
                f"\n⚠️ [SAYFA DÜZENİ HATASI: Sol kenar 3 cm, Sağ kenar 2.5 cm olmalı! "
                f"(Algılanan: Sol={l} cm, Sağ={r} cm)]"
            )
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            hata_sayisi += 1
            break

    # ── 2. İÇERİK DENETİMİ ───────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        upper_text = turkish_upper(text)
        style_name = para.style.name if para.style else ""

        # ── MOD GEÇİŞ KONTROLÜ ──────────────────────────────────────────

        # Ön sayfa bölümleri → tek satır modu
        if upper_text in ON_SAYFALAR_TEK:
            kapak_modu = False
            tek_satir_modu = True
            kaynakca_modu = False
            continue

        # Liste / dizin başlıkları (metin bazlı) → tek satır modu
        if any(upper_text.startswith(b) for b in LISTE_BASLIKLAR):
            kapak_modu = False
            tek_satir_modu = True
            kaynakca_modu = False
            continue

        # Kaynakça başlıkları → kaynakça modu (1.5 aralık)
        # Sadece Heading stili ile gelen KAYNAKLAR başlığı gerçek kaynakçadır.
        # TOC'taki "KAYNAKLAR\t96" gibi satırlar bu bloğa girmemeli.
        if (any(upper_text.startswith(b) for b in KAYNAKCA_BASLIKLAR)
                and len(text) < 35
                and style_name in HEADING_STILLER):
            kapak_modu = False
            kaynakca_modu = True
            tek_satir_modu = False
            continue

        # Heading stilli paragraflar → mod güncelle, kendisi denetlenmesin
        if style_name in HEADING_STILLER:
            kapak_modu = False
            # Heading içeriği liste başlığıysa tek satır moduna geç
            if any(upper_text.startswith(b) for b in LISTE_BASLIKLAR) or upper_text in ON_SAYFALAR_TEK:
                tek_satir_modu = True
                kaynakca_modu = False
            elif any(upper_text.startswith(b) for b in KAYNAKCA_BASLIKLAR) and len(text) < 35:
                kaynakca_modu = True
                tek_satir_modu = False
            else:
                tek_satir_modu = False
                kaynakca_modu = False
            continue

        # Numaralı bölüm başlıkları (1. GİRİŞ, 2.1 vb.) → ana metin modu
        if re.match(r'^\d+(\.\d+)*\.?\s+\S', text):
            kapak_modu = False
            tek_satir_modu = False
            kaynakca_modu = False
            continue

        # ── MUAF STİL KONTROLÜ ──────────────────────────────────────────
        if style_name in MUAF_STILLER:
            continue
        if "toc" in style_name.lower():
            continue

        # ── SATIR ARALIĞI DENETİMİ ──────────────────────────────────────
        effective_spacing = get_effective_spacing(para)
        hata_mesaji = ""

        if not kapak_modu:

            if tek_satir_modu or style_name in TEK_SATIR_STILLER:
                # TEK SATIR (1.0) bölgesi
                if effective_spacing is not None and effective_spacing > 1.15:
                    hata_mesaji += (
                        f" [HATA: Bu bölümde satır aralığı 1.0 (tek) olmalı! "
                        f"(Algılanan: {effective_spacing:.2f})]"
                    )

            elif kaynakca_modu or style_name in ANA_METIN_STILLER or effective_spacing is not None:
                # ANA METİN / KAYNAKÇA (1.5) bölgesi
                # "Şekil/Tablo/Resim/Not." ile başlayan satırlar muaf (Caption stili zaten muaf)
                if not text.startswith(("Şekil", "Tablo", "Resim", "Not.")):
                    if effective_spacing is not None and (
                        effective_spacing < 1.35 or effective_spacing > 1.65
                    ):
                        hata_mesaji += (
                            f" [HATA: Satır aralığı 1.5 olmalı! "
                            f"(Algılanan: {effective_spacing:.2f})]"
                        )

        # ── HATA İŞARETLE ───────────────────────────────────────────────
        if hata_mesaji:
            run = para.add_run(hata_mesaji)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            hata_sayisi += 1

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream, hata_sayisi


# ══════════════════════════════════════════════════════════════════════════════
# Streamlit arayüzü
# ══════════════════════════════════════════════════════════════════════════════

st.title("🎓 Balıkesir Üniversitesi Tez Denetim Sistemi (v30 – R11)")

st.markdown("""
**Uygulanan satır aralığı kuralları:**

| Bölüm | Kural |
|---|---|
| Ana metin, Kaynakça | 1,5 satır aralığı |
| Kabul ve Onay, Etik Beyan, Özet, Abstract, İçindekiler, Şekil/Tablo/Sembol/Kısaltma Listesi, Önsöz, Teşekkür | Tek satır (1,0) |
| Şekil/Tablo açıklamaları, alıntılar, dipnotlar | Tek satır — *stil tabanlı, otomatik muaf* |
""")

yuklenen_dosya = st.file_uploader("Tez dosyasını yükleyin (.docx)", type=["docx"])

if yuklenen_dosya and st.button("🔍 Denetimi Başlat"):
    with st.spinner("Belge taranıyor..."):
        islenmis_dosya, sayi = tez_denetle(yuklenen_dosya)

    if sayi == 0:
        st.success("✅ Herhangi bir hata tespit edilmedi.")
    else:
        st.warning(f"⚠️ Toplam **{sayi}** hata tespit edildi. Aşağıdan raporlanmış belgeyi indirin.")

    st.download_button(
        label="📥 Denetim Raporunu İndir",
        data=islenmis_dosya,
        file_name="TEZ_DENETIM_R11.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )