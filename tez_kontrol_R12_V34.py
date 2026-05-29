import streamlit as st
import re
import zipfile
from docx import Document
from io import BytesIO
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


# ══════════════════════════════════════════════════════════════════════════════
# Yardımcı fonksiyonlar
# ══════════════════════════════════════════════════════════════════════════════

def turkish_upper(text):
    replace_map = {"i": "İ", "ı": "I", "ç": "Ç", "ğ": "Ğ",
                   "ö": "Ö", "ş": "Ş", "ü": "Ü"}
    for lower, upper in replace_map.items():
        text = text.replace(lower, upper)
    return text.upper().strip()


# Türk alfabesi sıralama anahtarı
# Python'un Unicode karşılaştırması Türkçe harfleri (İ, Ğ, Ş, Ö, Ü, Ç) yanlış sıralar.
# Bu harita her harfi Türk alfabesindeki sıra numarasına çevirir.
_TR_ALFABE_SIRASI = {
    'A': 0,  'B': 1,  'C': 2,  'Ç': 3,  'D': 4,  'E': 5,  'F': 6,
    'G': 7,  'Ğ': 8,  'H': 9,  'I': 10, 'İ': 11, 'J': 12, 'K': 13,
    'L': 14, 'M': 15, 'N': 16, 'O': 17, 'Ö': 18, 'P': 19, 'R': 20,
    'S': 21, 'Ş': 22, 'T': 23, 'U': 24, 'Ü': 25, 'V': 26, 'Y': 27,
    'Z': 28,
}

def turkce_siralama_anahtari(metin):
    """Metni Türk alfabesine göre karşılaştırılabilir bir tuple'a dönüştürür."""
    return tuple(_TR_ALFABE_SIRASI.get(c, ord(c)) for c in metin)


def get_effective_spacing(para):
    """
    Paragrafın gerçek satır aralığını döndürür.
    Paragraf düzeyinde ayarlanmamışsa, stil hiyerarşisini yukarı doğru
    tarayarak kalıtılan (inherited) değeri bulur.
    """
    spacing = para.paragraph_format.line_spacing
    if spacing is not None:
        return spacing
    style_obj = para.style
    while style_obj is not None:
        s = style_obj.paragraph_format.line_spacing
        if s is not None:
            return s
        style_obj = style_obj.base_style
    return None


# ══════════════════════════════════════════════════════════════════════════════
# Sabitler
# ══════════════════════════════════════════════════════════════════════════════

# Asla satır aralığı denetimine tabi tutulmayan stiller
MUAF_STILLER = {
    "fbe_kapak_genel",
    "toc 1", "toc 2", "toc 3", "toc 4", "toc 7", "toc 8",
    "table of figures",
    "Caption",            # Şekil/Tablo/Resim başlıkları — zaten 1.0, muaf
    "annotation text",    # Tablo notları — zaten 1.0, muaf
    "fbe_metin_normal_kalın",  # "Sayfa" etiketi satırları
    "fbe_başlık1_sol",
    "List Paragraph",     # Madde/liste işareti paragrafları — kural kapsamı dışı
}

# Bölüm başlık stilleri — modu günceller, kendileri denetlenmez
HEADING_STILLER = {f"Heading {i}" for i in range(1, 10)}

# Bu stiller bulundukları bölümden bağımsız olarak TEK SATIR gerektirir
TEK_SATIR_STILLER = {
    "fbe_metin_sıkışık",   # Önsöz / Teşekkür gövde metni
    # NOT: "Normal (Web)" stili bazı tezlerde ana metin stili olarak kullanıldığından
    # burada sabit kodlanmaz; mod tabanlı kontrol ile ele alınır.
}

# Bu stiller ANA METİN kapsamında değerlendirilir (1.5 aralık)
ANA_METIN_STILLER = {
    "Tez Metni",
    "fbe_metin_normal",    # Kapak jüri listesi + bazı metin blokları
    "Default",             # Etik Beyan metni
    "Normal",              # Kaynakça + genel metin
    "Normal (Web)",        # Bazı tezlerde SONUÇLAR/ABSTRACT gibi ana metin stili
    "Bibliography",        # Kaynakça girişleri (Word'ün yerleşik kaynakça stili)
    "numarasiz",           # Numarasız liste / metin stili
    "Decimal Aligned",     # Ondalık hizalı tablo/metin stili
}

# Tek satır moduna geçiren bölüm başlıkları (metin olarak)
ON_SAYFALAR_TEK = {
    "ÖZET", "ABSTRACT", "ETİK BEYAN", "ÖNSÖZ",
    "TEŞEKKÜR", "KABUL VE ONAY", "DESTEKLEYICI KURULUŞ BEYAN SAYFASI",
}

# Tek satır moduna geçiren liste/dizin başlıkları (metin olarak)
LISTE_BASLIKLAR = {
    "İÇİNDEKİLER",
    "ŞEKİL LİSTESİ",
    "TABLO LİSTESİ",
    "RESİM LİSTESİ",
    "SEMBOL LİSTESİ",
    "KISALTMALAR LİSTESİ",
    "SİMGELER VE KISALTMALAR LİSTESİ",
    "SEMBOL VE KISALTMA LİSTESİ",
    # NOT: ÖZGEÇMİŞ bazı tezlerde 1.5 aralıkla yazıldığından buraya dahil edilmez.
}

# Kaynakça başlıkları
KAYNAKCA_BASLIKLAR = {
    "KAYNAKLAR", "KAYNAKÇA", "REFERENCES", "KAYNAKLAR DİZİNİ",
}


# ══════════════════════════════════════════════════════════════════════════════
# Sayfa numarası denetim yardımcıları
# ══════════════════════════════════════════════════════════════════════════════

def _footer_dosyalarini_analiz_et(docx_stream):
    """
    .docx içindeki tüm footer XML dosyalarını okuyarak her birinin
    sayfa numarası alanı, hizalama, font boyutu ve parantez/çizgi
    durumunu döndürür.
    """
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    if isinstance(docx_stream, str):
        zf = zipfile.ZipFile(docx_stream)
    else:
        docx_stream.seek(0)
        zf = zipfile.ZipFile(docx_stream)

    with zf:
        rels_xml = zf.read('word/_rels/document.xml.rels').decode('utf-8')
        rels_root = etree.fromstring(rels_xml.encode())
        rid_to_file = {
            r.get('Id'): r.get('Target')
            for r in rels_root
            if 'footer' in r.get('Type', '').lower()
        }

        footer_data = {}
        for rid, fname in rid_to_file.items():
            try:
                raw = zf.read(f'word/{fname}').decode('utf-8')
                tree = etree.fromstring(raw.encode())

                instr_els = tree.findall(f'.//{{{ns}}}instrText')
                instr_texts = [el.text for el in instr_els if el.text]
                has_page = any('PAGE' in t.upper() for t in instr_texts)

                jc_els = tree.findall(f'.//{{{ns}}}jc')
                alignment = jc_els[0].get(f'{{{ns}}}val') if jc_els else None

                sz_els = tree.findall(f'.//{{{ns}}}sz')
                font_sizes = sorted(set(
                    int(el.get(f'{{{ns}}}val', 0)) / 2
                    for el in sz_els
                    if el.get(f'{{{ns}}}val')
                ))

                # PAGE alanı DIŞINDA kalan metin (parantez/çizgi kontrolü)
                non_field_chars = []
                in_field = False
                for el in tree.iter():
                    tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                    if tag == 'fldChar':
                        ftype = el.get(f'{{{ns}}}fldCharType', '')
                        in_field = ftype in ('begin', 'separate')
                        if ftype == 'end':
                            in_field = False
                    elif tag == 't' and not in_field and el.text:
                        non_field_chars.append(el.text)

                non_field_text = ''.join(non_field_chars)
                has_bracket = any(c in non_field_text for c in ['(', ')', '-', '–', '—', '[', ']', '/'])

                footer_data[fname] = {
                    'has_page':       has_page,
                    'alignment':      alignment,
                    'font_sizes':     font_sizes,
                    'instr_texts':    instr_texts,
                    'has_bracket':    has_bracket,
                    'non_field_text': non_field_text,
                }
            except Exception as exc:
                footer_data[fname] = {'error': str(exc)}

    return rid_to_file, footer_data


def sayfa_numarasi_denetle(doc, docx_stream):
    """
    Kural listesi:
      1. Sayfa numarası altta ve ortalanmış olmalı (jc = center).
      2. Yazı boyutu 10 pt olmalı.
      3. Ön sayfalar → küçük harf Romen rakamı (lowerRoman).
      4. Giriş ve sonrası → Arap rakamı (decimal / None).
      5. Parantez veya çizgi kullanılmamalı.
    """
    hatalar = []

    rid_to_file, footer_data = _footer_dosyalarini_analiz_et(docx_stream)

    # Section → footer dosyası eşlemesi (yalnızca 'default' tipi)
    section_footer_map = []
    for si, section in enumerate(doc.sections):
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        fmt   = pgNumType.get(qn('w:fmt'))   if pgNumType is not None else None
        start = pgNumType.get(qn('w:start')) if pgNumType is not None else None

        # ── Kapak / ara geçiş section'larını muaf tut ──────────────────
        # Kriter: pgNumType.start tanımlı değilse bu section'da sayfa
        # numaralandırması başlatılmamıştır (kapak, ara geçiş sayfaları vb.).
        # Sayfa numarası olmaması bu durumlarda normaldir.
        if start is None:
            continue   # Numaralandırma başlatılmamış section — denetleme

        footer_refs_all = sectPr.findall(qn('w:footerReference'))

        for ref in footer_refs_all:
            ftype = ref.get(qn('w:type'))
            rid   = ref.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            fname = rid_to_file.get(rid)
            if fname and ftype == 'default':
                section_footer_map.append((si, fmt, start, fname))

    if not section_footer_map:
        hatalar.append("⚠️ [SAYFA NUMARASI: Footer tanımı bulunamadı — sayfa numarası eksik olabilir!]")
        return hatalar

    # Romen / Arap section varlığını TÜM section'lar üzerinden kontrol et
    # (footer referansı olmayan linked section'lar da sayılır)
    romen_var = False
    arap_var  = False
    for section in doc.sections:
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        fmt   = pgNumType.get(qn('w:fmt'))   if pgNumType is not None else None
        start = pgNumType.get(qn('w:start')) if pgNumType is not None else None
        if start is None:
            continue  # numaralandırma başlatılmamış — atla
        if fmt == 'lowerRoman':
            romen_var = True
        else:
            arap_var = True

    if not romen_var:
        hatalar.append(
            "⚠️ [SAYFA NUMARASI: Ön sayfa bölümü (Özet, İçindekiler vb.) için küçük harf Romen "
            "rakamı (i, ii, iii…) bulunamadı! İlgili section'da pgNumType=lowerRoman olmalı.]"
        )
    if not arap_var:
        hatalar.append(
            "⚠️ [SAYFA NUMARASI: Giriş ve sonrası için Arap rakamı (1, 2, 3…) bulunamadı! "
            "İlgili section'da pgNumType=decimal veya tanımsız olmalı.]"
        )

    # Her benzersiz footer dosyasını bir kez denetle
    denetlenen = set()
    for si, fmt, start, fname in section_footer_map:
        if fname in denetlenen:
            continue
        denetlenen.add(fname)

        data = footer_data.get(fname, {})
        if 'error' in data:
            hatalar.append(f"⚠️ [SAYFA NUMARASI: Section {si+1} footer okunamadı: {data['error']}]")
            continue

        etiket = f"Section {si + 1} ({'Romen/ön sayfalar' if fmt == 'lowerRoman' else 'Arap/ana metin'})"

        # Kural 1 — PAGE alanı var mı?
        if not data['has_page']:
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Footer'da sayfa numarası alanı (PAGE) bulunamadı!]"
            )
            continue

        # Kural 2 — Ortalanmış olmalı
        if data['alignment'] != 'center':
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası ortalanmış (alt-orta) olmalı! "
                f"(Algılanan hizalama: {data['alignment'] or 'sol/belirsiz'})]"
            )

        # Kural 3 — 10 pt olmalı
        sizes = data['font_sizes']
        if sizes and not all(abs(s - 10.0) < 0.6 for s in sizes):
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası boyutu 10 pt olmalı! "
                f"(Algılanan: {sizes} pt)]"
            )

        # Kural 4 — Parantez / çizgi olmamalı
        if data['has_bracket']:
            hatalar.append(
                f"⚠️ [SAYFA NUMARASI: {etiket} — Sayfa numarası yanında parantez veya çizgi "
                f"kullanılmamalı! (Algılanan: {repr(data['non_field_text'])})]"
            )

    return hatalar


# ══════════════════════════════════════════════════════════════════════════════
# Ana denetim fonksiyonu
# ══════════════════════════════════════════════════════════════════════════════

def tez_denetle(file_stream, progress_bar=None, progress_text=None):
    # Sayfa numarası denetimi için stream'i iki kez okumamız gerektiğinden kopyala
    file_stream.seek(0)
    docx_bytes = file_stream.read()

    doc = Document(BytesIO(docx_bytes))
    hata_sayisi = 0

    if progress_bar: progress_bar.progress(5)
    if progress_text: progress_text.text("Sayfa numaraları ve altbilgiler denetleniyor... %5")

    # ── 0. SAYFA NUMARASI DENETİMİ ───────────────────────────────────────
    sayfa_hatalari = sayfa_numarasi_denetle(doc, BytesIO(docx_bytes))
    if sayfa_hatalari:
        # Hataları belgenin ilk paragrafına ekle
        para = doc.paragraphs[0]
        for hata_metni in sayfa_hatalari:
            run = para.add_run("\n" + hata_metni)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
        hata_sayisi += len(sayfa_hatalari)

    # Durum değişkenleri
    kapak_modu = True
    tek_satir_modu = False
    kaynakca_modu = False
    son_kaynak = ""
    onceki_para_is_heading = False
    onceki_para_heading_text = ""
    bos_paragraf_sayaci = 0

    # ── 1. KENAR BOŞLUKLARI ───────────────────────────────────────────────
    for section in doc.sections:
        l = round(section.left_margin.cm, 1)
        r = round(section.right_margin.cm, 1)
        if l != 3.0 or r != 2.5:
            run = doc.paragraphs[0].add_run(
                f"\n⚠️ [SAYFA DÜZENİ HATASI: Sol kenar 3 cm, Sağ kenar 2.5 cm olmalı! "
                f"(Algılanan: Sol={l} cm, Sağ={r} cm)]"
            )
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            hata_sayisi += 1
            break

    # ── 2. İÇERİK DENETİMİ ───────────────────────────────────────────────
    total_paras = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        if progress_bar and total_paras > 0:
            # İşlem performansını düşürmemek için 20 dilimde bir veya son adımda güncelle
            if i % max(1, total_paras // 20) == 0 or i == total_paras - 1:
                percent = 5 + int((i + 1) / total_paras * 90)
                progress_bar.progress(percent)
                if progress_text: progress_text.text(f"Paragraflar inceleniyor... %{percent}")
        text = para.text.strip()
        if not text:
            bos_paragraf_sayaci += 1
            continue

        upper_text = turkish_upper(text)
        style_name = para.style.name if para.style else ""
        hata_mesaji = ""

        # Başlıklar sonrası boşluk kuralları kontrolü
        if onceki_para_is_heading:
            if onceki_para_heading_text == "ÖNSÖZ":
                if bos_paragraf_sayaci != 2:
                    hata_mesaji += f" [HATA: ÖNSÖZ başlığından sonra tam 2 adet boş paragraf bırakılmalı! (Bulunan: {bos_paragraf_sayaci})]"
            elif onceki_para_heading_text in ["KAYNAKLAR", "KAYNAKÇA", "REFERENCES"]:
                if bos_paragraf_sayaci != 1:
                    hata_mesaji += f" [HATA: {onceki_para_heading_text} başlığından sonra 1 adet boş paragraf bırakılmalı! (Bulunan: {bos_paragraf_sayaci})]"
            else:
                if bos_paragraf_sayaci > 0:
                    hata_mesaji += f" [HATA: Başlıktan sonra boş paragraf (enter) bırakılmamalıdır! (Bulunan: {bos_paragraf_sayaci})]"
        
        bos_paragraf_sayaci = 0 # Sıfırla

        # ── TOC (İÇİNDEKİLER / LİSTELER) KONTROLÜ ───────────────────────
        is_toc_entry = "toc" in style_name.lower() or bool(re.search(r'\.{3,}\s*[a-zA-Z0-9ivxlcdmIVXLCDM]+$', text.strip()))
        if not is_toc_entry and onceki_para_heading_text in LISTE_BASLIKLAR:
            # Nokta yok ama bölüm İçindekiler/Listeler ise ve sağda sayfa no formatı varsa
            if re.search(r'\s+[a-zA-Z0-9ivxlcdmIVXLCDM]+$', text.strip()) and len(text) < 150:
                is_toc_entry = True

        if is_toc_entry:
            onceki_para_is_heading = False
            # TOC girdileri için başlık ve iki yana yaslı kurallarını atlıyoruz.
            continue

        # ── MOD GEÇİŞ KONTROLÜ ──────────────────────────────────────────

        # Ön sayfa bölümleri → tek satır modu
        if upper_text in ON_SAYFALAR_TEK:
            kapak_modu = False
            tek_satir_modu = True
            kaynakca_modu = False
            son_kaynak = ""
            onceki_para_is_heading = True
            onceki_para_heading_text = upper_text
            continue

        # Liste / dizin başlıkları (metin bazlı) → tek satır modu
        if any(upper_text.startswith(b) for b in LISTE_BASLIKLAR):
            kapak_modu = False
            tek_satir_modu = True
            kaynakca_modu = False
            son_kaynak = ""
            onceki_para_is_heading = True
            onceki_para_heading_text = next((b for b in LISTE_BASLIKLAR if upper_text.startswith(b)), upper_text)
            continue

        # Kaynakça başlıkları → kaynakça modu (1.5 aralık)
        if (any(upper_text.startswith(b) for b in KAYNAKCA_BASLIKLAR)
                and len(text) < 35
                and style_name in HEADING_STILLER):
            kapak_modu = False
            kaynakca_modu = True
            tek_satir_modu = False
            son_kaynak = ""
            onceki_para_is_heading = True
            onceki_para_heading_text = next((b for b in KAYNAKCA_BASLIKLAR if upper_text.startswith(b)), upper_text)
            continue

        # Heading stilli paragraflar → mod güncelle, kendisi denetlenmesin
        if style_name in HEADING_STILLER:
            kapak_modu = False
            onceki_para_is_heading = True
            onceki_para_heading_text = upper_text

            # Keep with next kontrolü
            # Şimdilik bu kontrol pasife alındı
            # if not para.paragraph_format.keep_with_next:
            #     hata_mesaji += " [HATA: Başlığın son sayfada tek kalmaması için 'Sonrakiyle Birlikte Tut' (Keep with next) özelliği açık olmalıdır!]"
                
            # İki yana yaslı kontrolü
            if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                hata_mesaji += " [HATA: Tüm başlıklar iki yana yaslı (Justify) olmalıdır!]"

            # 3. ve 4. seviye başlıkların tamamı büyük harf olmamalı
            if style_name in ["Heading 3", "Heading 4"] and text.isupper():
                hata_mesaji += " [HATA: 3. ve 4. seviye başlıklarda sadece kelimelerin ilk harfi büyük olmalıdır (Tamamı büyük yazılmış)!]"

            # Hata varsa ekleyelim çünkü continue ile atlayacağız
            if hata_mesaji:
                run = para.add_run(hata_mesaji)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                hata_sayisi += 1

            # Heading içeriği liste başlığıysa tek satır moduna geç
            if any(upper_text.startswith(b) for b in LISTE_BASLIKLAR) or upper_text in ON_SAYFALAR_TEK:
                tek_satir_modu = True
                kaynakca_modu = False
            elif any(upper_text.startswith(b) for b in KAYNAKCA_BASLIKLAR) and len(text) < 35:
                kaynakca_modu = True
                tek_satir_modu = False
                son_kaynak = ""
            else:
                tek_satir_modu = False
                kaynakca_modu = False
            continue

        # Eğer buradaysak paragraf başlık değildir
        onceki_para_is_heading = False

        # Numaralı bölüm başlıkları (1. GİRİŞ, 2.1 vb.) → ana metin modu
        if re.match(r'^\d+(\.\d+)*\.?\s+\S', text):
            kapak_modu = False
            tek_satir_modu = False
            kaynakca_modu = False
            continue

        # ── YAZI TİPİ KONTROLÜ ──────────────────────────────────────────
        has_invalid_font = False
        has_invalid_size = False
        for run in para.runs:
            if run.text.strip():
                if run.font.name and run.font.name != "Times New Roman":
                    has_invalid_font = True
                
                # Boyut kontrolü (Sadece başlık/muaf olmayan ana metinler için)
                if not kapak_modu and style_name not in HEADING_STILLER and style_name not in MUAF_STILLER:
                    size = run.font.size
                    if size is None and para.style and para.style.font:
                        size = para.style.font.size
                    if size and size.pt != 12.0:
                        has_invalid_size = True

        if has_invalid_font:
            hata_mesaji += " [HATA: Yazı tipi Times New Roman olmalıdır!]"
        if has_invalid_size:
            hata_mesaji += " [HATA: Yazı boyutu 12 pt olmalıdır!]"

        # ── ŞEKİL/TABLO BOLD KONTROLÜ ──────────────────────────────────
        if text.startswith(("Şekil", "Tablo", "Resim")):
            for run in para.runs:
                # İçerisinde en az bir harf veya rakam olan ilk run'ı bul (alan kodlarını atlar)
                if re.search(r'[a-zA-Z0-9çğıöşüÇĞİÖŞÜ]', run.text):
                    is_bold = run.bold
                    if is_bold is None:
                        # Stil kalıtımına bak
                        style = para.style
                        while style is not None:
                            if style.font.bold is not None:
                                is_bold = style.font.bold
                                break
                            style = style.base_style
                            
                    if not is_bold:
                        hata_mesaji += " [HATA: Şekil/Tablo/Resim numaralandırma kısmı kalın (bold) olmalıdır!]"
                    break

        # ── MUAF STİL KONTROLÜ ──────────────────────────────────────────
        if style_name in MUAF_STILLER:
            if hata_mesaji:
                run = para.add_run(hata_mesaji)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                hata_sayisi += 1
            continue
        # ── HİZALAMA VE GİRİNTİ KONTROLÜ (Sadece ana metin/kaynakça) ─────
        if not kapak_modu:
            align = para.paragraph_format.alignment
            
            # ÖZET ve ABSTRACT içindeki başlık bloğunu (ortalanmış, kısa veya büyük harfli) muaf tut
            is_title_block = False
            if align == WD_ALIGN_PARAGRAPH.CENTER and onceki_para_heading_text in ["ÖZET", "ABSTRACT"]:
                if text == upper_text or re.search(r'(TEZİ|THESIS|DANIŞMAN|SUPERVISOR|ÜNİVERSİTE|UNIVERSITY|ENSTİTÜ|INSTITUTE|ANABİLİM|DEPARTMENT)', upper_text) or len(text) < 150:
                    is_title_block = True

            # Tüm yazılar iki yana yaslı olmalı
            if align not in [WD_ALIGN_PARAGRAPH.JUSTIFY, None] and not is_title_block:
                hata_mesaji += " [HATA: Paragraf iki yana yaslı olmalıdır!]"
            
            # Paragraf başı (girinti) olmamalı
            indent = para.paragraph_format.first_line_indent
            if indent is not None and indent.pt > 0:
                hata_mesaji += " [HATA: Paragraf başı (girinti) yapılmamalıdır!]"

        # ── PARAGRAF ÖNCESİ/SONRASI BOŞLUK KONTROLÜ (0 nk) ───────────────
        if not kapak_modu and style_name not in HEADING_STILLER:
            before = para.paragraph_format.space_before
            if before is not None and before.pt > 0:
                hata_mesaji += " [HATA: Paragraf öncesi boşluk (Space Before) 0 nk olmalıdır!]"
            after = para.paragraph_format.space_after
            if after is not None and after.pt > 0:
                hata_mesaji += " [HATA: Paragraf sonrası boşluk (Space After) 0 nk olmalıdır!]"

        # ── SATIR ARALIĞI DENETİMİ ──────────────────────────────────────
        effective_spacing = get_effective_spacing(para)

        if not kapak_modu:

            if tek_satir_modu or style_name in TEK_SATIR_STILLER:
                # TEK SATIR (1.0) bölgesi
                if effective_spacing is not None and effective_spacing > 1.15:
                    hata_mesaji += (
                        f" [HATA: Bu bölümde satır aralığı 1.0 (tek) olmalı! "
                        f"(Algılanan: {effective_spacing:.2f})]"
                    )

            elif kaynakca_modu or style_name in ANA_METIN_STILLER or effective_spacing is not None:
                # ANA METİN / KAYNAKÇA (1.5) bölgesi
                # "Şekil/Tablo/Resim/Not." ile başlayan satırlar muaf
                if not text.startswith(("Şekil", "Tablo", "Resim", "Not.")):
                    if effective_spacing is not None and (
                        effective_spacing < 1.35 or effective_spacing > 1.65
                    ):
                        hata_mesaji += (
                            f" [HATA: Satır aralığı 1.5 olmalı! "
                            f"(Algılanan: {effective_spacing:.2f})]"
                        )

        # ── ALFABETİK SIRA VE ASILI GİRİNTİ (KAYNAKÇA) ──────────────────
        KAYNAK_DISI_ONEKLER = (
            "KAYNAKLAR", "KAYNAKÇA", "REFERENCES", "EKLER", "ÖZGEÇMİŞ", "SAYFA"
        )
        kaynak_disi = any(upper_text.startswith(k) for k in KAYNAK_DISI_ONEKLER)

        if kaynakca_modu and len(text) > 3 and not kaynak_disi:
            # Asılı Girinti Kontrolü (Hanging Indent)
            indent = para.paragraph_format.first_line_indent
            if indent is None or indent.pt >= 0:
                hata_mesaji += " [HATA: Kaynakça öğesi asılı girinti (hanging indent) ile yazılmalıdır!]"
                
            # Satırları Birlikte Tut Kontrolü
            if not para.paragraph_format.keep_together:
                hata_mesaji += " [HATA: Kaynakça maddesinin bölünmemesi için 'Satırları Birlikte Tut' (Keep lines together) açık olmalıdır!]"

            pure = re.sub(r'^\[\d+\]\s*', '', text).strip()
            current_val = turkish_upper(pure)
            if current_val and son_kaynak:
                # Türk alfabesi sırasına göre karşılaştır
                # Şimdilik alfabetik sıralama pasife alındı
                # if turkce_siralama_anahtari(current_val) < turkce_siralama_anahtari(son_kaynak):
                #     hata_mesaji += " [HATA: Kaynakça alfabetik sırayı bozuyor!]"
                pass
            son_kaynak = current_val

        # ── HATA İŞARETLE ───────────────────────────────────────────────
        if hata_mesaji:
            run = para.add_run(hata_mesaji)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True
            hata_sayisi += 1

    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream, hata_sayisi


# ══════════════════════════════════════════════════════════════════════════════
# Streamlit arayüzü
# ══════════════════════════════════════════════════════════════════════════════

st.title("🎓 Balıkesir Üniversitesi Tez Denetim Sistemi (v34 – R12)")

st.markdown("""
**Uygulanan satır aralığı kuralları:**

| Bölüm | Kural |
|---|---|
| Ana metin, Kaynakça | 1,5 satır aralığı |
| Kabul ve Onay, Etik Beyan, Özet, Abstract, İçindekiler, Şekil/Tablo/Sembol/Kısaltma Listesi, Önsöz, Teşekkür | Tek satır (1,0) |
| Şekil/Tablo açıklamaları, alıntılar, dipnotlar | Tek satır — *stil tabanlı, otomatik muaf* |
""")

yuklenen_dosya = st.file_uploader("Tez dosyasını yükleyin (.docx)", type=["docx"])

if yuklenen_dosya and st.button("🔍 Denetimi Başlat"):
    progress_bar = st.progress(0)
    progress_text = st.empty()
    progress_text.text("Belge taranmaya başlanıyor... %0")
    
    with st.spinner("Arka planda denetim devam ediyor... Lütfen bekleyin."):
        islenmis_dosya, sayi = tez_denetle(yuklenen_dosya, progress_bar, progress_text)
        
    progress_bar.progress(100)
    progress_text.text("Denetim tamamlandı! %100")

    if sayi == 0:
        st.success("✅ Herhangi bir hata tespit edilmedi.")
    else:
        st.warning(f"⚠️ Toplam **{sayi}** hata tespit edildi. Aşağıdan raporlanmış belgeyi indirin.")

    st.download_button(
        label="📥 Denetim Raporunu İndir",
        data=islenmis_dosya,
        file_name="TEZ_DENETIM_R10.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
